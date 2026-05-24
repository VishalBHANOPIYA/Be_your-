import os
import logging
import pdfplumber
import docx
import re

logger = logging.getLogger(__name__)

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(pdf_path):
        text = ""
        # Layer 1: pdfplumber (Primary)
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
            
        text = text.strip()
        if len(text) > 20:
            return text
            
        # Layer 2: pypdfium2 (Secondary Fallback)
        logger.info("pdfplumber returned empty/too-short text. Trying pypdfium2 fallback...")
        pypdf_text = ""
        try:
            import pypdfium2 as pdfium
            doc = pdfium.PdfDocument(pdf_path)
            for page in doc:
                textpage = page.get_textpage()
                extracted = textpage.get_text_bounded()
                if extracted:
                    pypdf_text += extracted + "\n"
        except Exception as e:
            logger.warning(f"pypdfium2 extraction failed: {e}")
            
        pypdf_text = pypdf_text.strip()
        if len(pypdf_text) > 20:
            return pypdf_text
            
        # Layer 3: pdfminer (Tertiary Fallback)
        logger.info("pypdfium2 returned empty/too-short text. Trying pdfminer fallback...")
        pdfminer_text = ""
        try:
            from pdfminer.high_level import extract_text
            extracted = extract_text(pdf_path)
            if extracted:
                pdfminer_text = extracted
        except Exception as e:
            logger.warning(f"pdfminer extraction failed: {e}")
            
        pdfminer_text = pdfminer_text.strip()
        if len(pdfminer_text) > 20:
            return pdfminer_text
            
        return text or pypdf_text or pdfminer_text or ""


    @staticmethod
    def extract_text_from_docx(docx_path):
        text = ""
        try:
            doc = docx.Document(docx_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
        return text

    @staticmethod
    def extract_text(file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return ResumeParser.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return ResumeParser.extract_text_from_docx(file_path)
        return ""

    @staticmethod
    def clean_text(text):
        # Remove extra whitespaces and special characters
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\x00-\x7f]', r'', text) # Remove non-ascii
        return text.strip()
